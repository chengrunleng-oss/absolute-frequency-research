from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT
FIG = ROOT / "assets" / "paper_figures"

INK = "20313F"
BLUE = "1F5D78"
TEAL = "2B7A78"
MUTED = "61717E"
LIGHT = "EAF1F4"
GOLD = "B78B2E"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Microsoft YaHei", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.34)
    section.footer_distance = Inches(0.34)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.22

    for name, size, before, after, color in (
        ("Title", 27, 0, 8, INK),
        ("Subtitle", 12, 0, 14, MUTED),
        ("Heading 1", 17, 13, 6, BLUE),
        ("Heading 2", 12.5, 9, 4, TEAL),
        ("Heading 3", 11, 7, 3, INK),
    ):
        s = styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.bold = name != "Subtitle"
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Microsoft YaHei"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.keep_with_next = True


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("论文内容汇报  |  Lu+ 光频标绝对频率测量")
    set_font(r, size=8, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("arXiv:2502.10004v3  |  ")
    set_font(r, size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text.upper())
    set_font(r, size=9, bold=True, color=GOLD)


def add_lead_box(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.86)
    cell = table.cell(0, 0)
    cell.width = Inches(6.86)
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, top=150, bottom=150, start=190, end=190)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_font(r, size=10, bold=True, color=BLUE)
    r = p.add_run(text)
    set_font(r, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.23)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_figure(doc, n, title, why, reading, takeaway, width):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("看图目的：")
    set_font(r, size=9.5, bold=True, color=TEAL)
    r = p.add_run(why)
    set_font(r, size=9.5, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(0)
    shape = p.add_run().add_picture(str(FIG / f"figure_{n}.png"), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("title", f"图 {n} {title}")
    doc_pr.set("descr", f"原论文 Figure {n}：{title}。{takeaway}")

    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(f"图 {n}  {title}（原论文 Figure {n}）")

    if reading:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("读图要点：")
        set_font(r, size=9.5, bold=True, color=BLUE)
        r = p.add_run(reading)
        set_font(r, size=9.5, color=INK)
    add_lead_box(doc, "本图结论", takeaway)


def add_summary_table(doc):
    rows = [
        ("对象", "176Lu+ 单离子 1S0 -> 3D1 光学跃迁"),
        ("目标", "将 Lu+ 光钟频率溯源到国际单位制 SI 秒"),
        ("测量链", "Lu+ -> 光频梳 -> 氢钟 HM -> UTC(USNO) -> TAI -> SI"),
        ("数据周期", "2021 年 3-5 月，覆盖 Circular T 399-401"),
        ("最终结果", "353 638 794 073 800.35(33) Hz"),
        ("相对不确定度", "9.2 x 10^-16"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    widths = [1.35, 5.5]
    for i, (k, v) in enumerate(rows):
        for j, text in enumerate((k, v)):
            cell = table.cell(i, j)
            cell.width = Inches(widths[j])
            set_cell_margins(cell)
            if j == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_font(r, size=9.5, bold=(j == 0), color=BLUE if j == 0 else INK)


def add_uncertainty_budget_table(doc):
    rows = [
        ("统计分量 u_A", "", "", "", ""),
        ("Lu/TAI", "9.9", "13", "19", "7.3"),
        ("TAI interpolation", "1.7", "2.9", "6.3", "1.5"),
        ("TAI/SI", "0.7", "0.7", "0.9", "0.4"),
        ("统计合计", "10", "13", "20", "7.4"),
        ("系统分量 u_B", "", "", "", ""),
        ("HM drift", "2.0", "2.0", "2.0", "2.0"),
        ("HM diurnal", "5.0", "5.0", "5.0", "5.0"),
        ("Lu systematics", "0.1", "0.1", "0.1", "0.1"),
        ("Gravitational shift", "0.4", "0.4", "0.4", "0.4"),
        ("PSFS", "0.9", "0.8", "0.9", "0.9"),
        ("系统合计", "5.5", "5.5", "5.5", "5.5"),
        ("Lu/SI 总不确定度", "11", "14", "20", "9.2"),
    ]
    headers = ("不确定度来源", "Cir. T 399", "Cir. T 400", "Cir. T 401", "合并结果")
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    widths = [2.55, 1.05, 1.05, 1.05, 1.16]
    for j, text in enumerate(headers):
        cell = table.cell(0, j)
        cell.width = Inches(widths[j])
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell, top=90, bottom=90, start=80, end=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=8.2, bold=True, color="FFFFFF")
    set_repeat_table_header(table.rows[0])

    for label, *values in rows:
        cells = table.add_row().cells
        is_group = label in ("统计分量 u_A", "系统分量 u_B")
        is_total = label in ("统计合计", "系统合计", "Lu/SI 总不确定度")
        for j, text in enumerate((label, *values)):
            cells[j].width = Inches(widths[j])
            set_cell_margins(cells[j], top=65, bottom=65, start=80, end=80)
            if is_group:
                set_cell_shading(cells[j], LIGHT)
            elif label == "Lu/SI 总不确定度":
                set_cell_shading(cells[j], "DCEBED")
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_font(r, size=8.2, bold=(is_group or is_total), color=BLUE if is_group else INK)


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_font(r, name="Cambria Math", size=10.5, italic=True, color=INK)


def build():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)

    add_kicker(doc, "Research briefing")
    p = doc.add_paragraph(style="Title")
    p.add_run("Lu+ 光频标绝对频率测量\n内容汇报")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("基于论文 Absolute frequency measurement of a Lu+ (3D1) optical frequency standard via link to international atomic time")

    add_lead_box(doc, "核心结论", "作者首次完成 176Lu+ (3D1) 光频标的绝对频率测量，并通过本地氢钟、GPS 精密单点定位、UTC(USNO) 与 TAI 将结果连接到 SI 秒。")
    doc.add_heading("一、论文概览", level=1)
    add_summary_table(doc)
    doc.add_heading("汇报逻辑", level=2)
    add_bullets(doc, [
        "先说明 Lu+ 光钟的物理定义与询问方法（图 1）。",
        "再展开从实验室到 SI 秒的完整溯源链（图 2）。",
        "随后建立氢钟噪声模型，并说明原始数据如何校正、分箱（图 3-4）。",
        "最后评估光钟停机带来的插值不确定度，合并三个月结果（图 5-6）。",
    ])
    p = doc.add_paragraph()
    r = p.add_run("汇报重点：")
    set_font(r, bold=True, color=BLUE)
    p.add_run("论文的难点不只是测得 353 THz 量级的光频率，而是把不同设备、地点和时间窗口中的频率比与不确定度严密连接起来。")

    doc.add_page_break()
    doc.add_heading("二、光钟如何产生可测的标准频率", level=1)
    p = doc.add_paragraph("论文使用单个 176Lu+ 离子的 1S0 -> 3D1 跃迁作为钟跃迁。通过光学脉冲和两路微波脉冲在三个超精细态之间进行相干转移，再对三个跃迁频率做超精细平均，降低对外磁场等扰动的敏感性。")
    add_figure(doc, 1, "Lu+ 能级结构、耦合跃迁与 hyper-Ramsey 询问时序",
               "回答“光钟频率由什么物理过程定义”。",
               "(a) 给出冷却、探测与钟跃迁相关能级；(b) 将钟激光与两路微波连接到三个 3D1 超精细态；(c) 展示对称的光学-微波脉冲序列。",
               "Lu+ 标准频率不是单条谱线的简单读数，而是由相干询问实现的超精细平均频率；对称 hyper-Ramsey 序列同时用于抑制探测光引起的频移。", 5.45)

    doc.add_page_break()
    doc.add_heading("三、从 Lu+ 光钟连接到 SI 秒", level=1)
    p = doc.add_paragraph("光钟只在实验运行时提供高精度光学频率。作者用光频梳把 353 THz 的钟激光转换到可与本地氢钟比较的射频信号；氢钟作为连续运行的“飞轮”，再经 GNSS 链路与 UTC(USNO) 比较，最后通过 BIPM 的 Circular T 接入 TAI 和 SI 秒。")
    add_figure(doc, 2, "Lu+ 光钟至国际原子时和 SI 秒的远程频率比较链",
               "回答“实验室中的光学频率如何获得 SI 溯源”。",
               "左侧是实验信号链：Lu+、钟激光、光频梳、SDR 与氢钟；中间和右侧是时间传递与计量链：GPS/PPP、UTC(USNO)、TAI 和 SI 秒。",
               "整篇论文可概括为五段频率比的连接：Lu/HM、HM/UTC(USNO)、UTC(USNO)/TAI、时间窗口插值以及 TAI/SI。每一段都必须给出对应的不确定度。", 6.65)

    doc.add_page_break()
    doc.add_heading("四、氢钟的稳定度与噪声模型", level=1)
    p = doc.add_paragraph("氢钟 HM 连续运行，但其频率会漂移并包含多种噪声。作者分别使用 Lu-HM 的连续比较和 HM-UTC(USNO) 的 PPP 比较，计算 Hadamard deviation，并据此构建后续停机插值模拟所需的 HM 噪声模型。")
    add_figure(doc, 3, "HM 相对 Lu+ 与 UTC(USNO) 的稳定度及模拟噪声模型",
               "回答“氢钟在不同平均时间上有多稳定，以及应如何模拟”。",
               "蓝点描述短时间 Lu-HM 比较，橙点描述长时间 PPP 链路；红线代表不同噪声分量，黑线为合成模型。短时段接近白频率噪声，中间出现平台，长时间约在 1.3 x 10^-15 附近。",
               "图 3 不是独立的性能展示，而是图 5 蒙特卡洛插值评估的模型输入。采用 Hadamard deviation 可以减弱线性频率漂移对稳定度估计的干扰。", 5.25)

    doc.add_page_break()
    doc.add_heading("五、原始观测如何变成 5 天频率结果", level=1)
    p = doc.add_paragraph("作者用约 80 天 HM-UTC(USNO) 数据拟合氢钟频率偏置、线性漂移和老化，再从 Lu-HM 与 HM-UTC(USNO) 数据中扣除共同趋势。随后将结果按 Circular T 的 5 天网格分箱，形成可与 TAI 链接的九个观测点。")
    add_figure(doc, 4, "氢钟漂移拟合、残差校正与 5 天 Lu/TAI 结果",
               "回答“分散且间歇的实验数据如何被整理成可用于绝对频率测量的结果”。",
               "(a) 蓝点为 Lu-HM 间歇测量，橙点为 HM-UTC(USNO) 六小时平均，黑线为二次拟合；(b) 展示扣除偏置和漂移后的残差；(c) 给出九个 5 天结果，黑色外误差棒为总统计不确定度，红色内误差棒仅表示 HM dead-time 插值贡献。",
               "这张图体现论文的主分析流程：先用连续 PPP 数据掌握氢钟长期行为，再校正光钟运行片段，最后与 TAI 的 5 天报告窗口对齐。约化卡方 0.67 表明九个点的散布与给定统计不确定度相容。", 4.75)

    doc.add_page_break()
    doc.add_heading("六、光钟停机造成的插值不确定度", level=1)
    p = doc.add_paragraph("在每个 5 天窗口内，Lu+ 光钟的实际运行比例仅约 2%-37%。因此，光钟运行时段的 HM 平均频率并不一定等于完整 5 天的 HM 平均频率。作者利用图 3 的噪声模型和真实 uptime 分布进行蒙特卡洛模拟，估计这种 dead time 带来的不确定度。")
    add_figure(doc, 5, "5 天窗口内光钟 uptime 与 HM 插值不确定度的关系",
               "回答“光钟少运行一段时间，会给最终频率增加多少不确定度”。",
               "横轴是 5 天内光钟运行比例，纵轴是 HM 插值不确定度。黑点对应实验中的真实运行分布；灰带和两条彩色曲线表示不同模拟调度情形。总体趋势是 uptime 越高，插值不确定度越小。",
               "提高 uptime 不仅增加数据量，更重要的是减少用 HM 跨越停机区间时的外推风险。相同 uptime 下，运行片段在时间上的分布也会影响不确定度。", 5.15)

    doc.add_page_break()
    doc.add_heading("七、不确定度预算表及分析", level=1)
    p = doc.add_paragraph("来源说明：下表为原论文 Table 3 的汇报版。所有数值均为相对标准不确定度，单位为 10^-16；Cir. T 399-401 是三期月度估计，合并结果按论文的相关性模型计算。")
    add_uncertainty_budget_table(doc)
    doc.add_heading("7.1 合成方法", level=2)
    p = doc.add_paragraph("【基础补充】相互独立的标准不确定度分量采用方差相加，即平方和开根号。若分量相关，则还必须加入协方差项，不能直接使用下面的简式。")
    add_equation(doc, "u_c = sqrt(sum(u_i^2))")
    p = doc.add_paragraph("合并结果的统计分量、系统分量和总不确定度分别为：")
    add_equation(doc, "u_A = sqrt(7.3^2 + 1.5^2 + 0.4^2) = 7.46")
    p = doc.add_paragraph("表中的 7.3、1.5、0.4 已经舍入，因此显示值复算得到 7.46；论文使用未舍入的内部数值合成后报告为 7.4。这是显示精度差异，不是合成规则不同。")
    add_equation(doc, "u_B = sqrt(2.0^2 + 5.0^2 + 0.1^2 + 0.4^2 + 0.9^2) = 5.5")
    add_equation(doc, "u_Lu/SI = sqrt(7.4^2 + 5.5^2) = 9.2")
    doc.add_heading("7.2 预算主导项", level=2)
    add_bullets(doc, [
        "统计侧主要由 Lu/TAI = 7.3 决定；它明显高于 TAI interpolation = 1.5 和 TAI/SI = 0.4。",
        "系统侧主要由 HM diurnal = 5.0 决定，反映光钟运行时段偏向白天时潜在的氢钟昼夜频率偏差。",
        "Cir. T 401 的 Lu/TAI = 19、TAI interpolation = 6.3，总不确定度达到 20，因此最终权重最低。",
        "Lu systematics = 0.1、引力红移 = 0.4，说明当前瓶颈主要在 HM 和远程时间链路，而不是 Lu+ 跃迁本体。",
    ])
    doc.add_heading("7.3 相关性与改进方向", level=2)
    p = doc.add_paragraph("HM drift、HM diurnal、Lu 系统频移、引力红移以及大部分 PSFS 系统分量在三个月中共享设备、模型或校准来源，因此不能按三个独立样本平均。统计项可随独立数据增加而下降，但合并后的系统合计仍为 5.5，而不是 5.5/sqrt(3)。")
    add_lead_box(doc, "改进优先级", "优先提高光钟 uptime、改善 HM 昼夜效应评估和远程频率传递；单独继续降低已经很小的 Lu+ 本体系统不确定度，对最终 9.2 x 10^-16 的改善有限。")

    doc.add_page_break()
    doc.add_heading("八、三个月结果与最终绝对频率", level=1)
    p = doc.add_paragraph("九个 5 天结果按 Circular T 399、400、401 三个月度窗口合并，再结合 TAI 相对 SI 秒的校准结果。跨月共同的系统不确定度不能被当作三次独立测量而简单平均掉，作者据此处理相关项并计算最终加权均值。")
    add_figure(doc, 6, "Circular T 399-401 月度 Lu/SI 结果与最终加权均值",
               "回答“各月结果是否一致，以及最终频率是多少”。",
               "三个黑点表示各月相对加权均值的频率偏差；红色误差棒仅为统计分量，黑色误差棒为总不确定度；灰带表示最终加权结果的相对不确定度范围。第三个月总误差最大。",
               "三个月度估计在不确定度范围内一致。最终得到 f_Lu = 353 638 794 073 800.35(33) Hz，相对不确定度为 9.2 x 10^-16。", 5.25)

    doc.add_heading("九、总结与评价", level=1)
    add_bullets(doc, [
        "主要贡献：首次给出 176Lu+ (3D1) 光频标的 SI 可溯源绝对频率，为未来形成推荐频率值提供关键实验依据。",
        "方法特点：把单离子光钟、光频梳、本地氢钟、GNSS PPP、UTC(USNO)、TAI 和 SI 秒组织成完整计量链。",
        "主要限制：结果仍明显受远程时间传递、氢钟停机插值、TAI 插值和引力红移评估影响；并非完全由 Lu+ 光钟本体性能决定。",
        "改进方向：提高光钟连续运行率，采用 PPP-AR 或光纤链路，改善氢钟/钟组稳定度，并进一步降低系统频移与高度测量不确定度。",
    ])
    add_lead_box(doc, "一句话总结", "论文证明了 Lu+ 光钟不仅可以在实验室内达到很高精度，也能够通过现实的国际时间链路给出可信的 SI 绝对频率。")

    doc.add_heading("资料说明", level=2)
    p = doc.add_paragraph("本文档依据 arXiv:2502.10004v3 及仓库中的论文解读资料整理。图 1-6 均直接取自原论文，中文说明为汇报性概括。论文未公开完整 Lu/HM 时间序列、本地 RINEX、uptime mask 和部分系统评估原始记录，因此外部读者可以复核方法与公开表格，但不能仅凭 PDF 独立重算最终中心值。")
    p.paragraph_format.space_after = Pt(0)

    path = OUT / "Lu光频标绝对频率测量_内容汇报.docx"
    doc.save(path)
    print(path)


if __name__ == "__main__":
    build()
